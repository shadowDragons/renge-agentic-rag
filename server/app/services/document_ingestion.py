from pathlib import Path
from shutil import copyfileobj
import subprocess
from uuid import uuid4

from fastapi import UploadFile
from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.orm import Session as DBSession

from app.core.config import get_settings
from app.db.session import SessionLocal
from app.integrations.llamaindex_ingestion import run_document_ingestion_pipeline
from app.integrations.qdrant_store import QdrantChunkStore
from app.models import Document, DocumentChunk, Job


SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".pdf",
    ".doc",
    ".docx",
}
SUPPORTED_DOCUMENT_FORMAT_LABEL = "TXT、Markdown、CSV、JSON、YAML、XML、HTML、PDF、DOC、DOCX"


class DocumentIngestionService:
    def __init__(self, db: DBSession):
        self.db = db
        self.settings = get_settings()

    def create_upload_task(
        self,
        knowledge_base_id: str,
        upload_file: UploadFile,
    ) -> tuple[Document, Job]:
        safe_name = Path(upload_file.filename or "未命名文件").name or "未命名文件"
        validate_supported_document_filename(safe_name)
        document_id = str(uuid4())
        job_id = str(uuid4())
        upload_dir = Path(self.settings.storage_root) / "uploads" / knowledge_base_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        storage_path = upload_dir / f"{document_id}__{safe_name}"

        with storage_path.open("wb") as target:
            copyfileobj(upload_file.file, target)

        document = Document(
            document_id=document_id,
            knowledge_base_id=knowledge_base_id,
            file_name=safe_name,
            file_path=str(storage_path),
            mime_type=upload_file.content_type or "application/octet-stream",
            status="processing",
        )
        job = Job(
            job_id=job_id,
            job_type="document_ingestion",
            target_id=document_id,
            status="pending",
            progress=0.0,
            error_message="",
        )
        self.db.add(document)
        self.db.add(job)
        self.db.commit()
        self.db.refresh(document)
        self.db.refresh(job)
        return document, job

    def retry_job(
        self,
        job_id: str,
    ) -> tuple[Document, Job]:
        job = self.db.scalar(select(Job).where(Job.job_id == job_id))
        if job is None:
            raise DocumentIngestionStateError("任务不存在。", status_code=404)
        return self._retry_document_job(job)

    def retry_jobs(
        self,
        *,
        job_ids: list[str] | None = None,
        limit: int = 20,
    ) -> tuple[list[tuple[Document, Job]], list[str]]:
        if limit <= 0:
            raise DocumentIngestionStateError("批量重试 limit 必须大于 0。", status_code=422)

        retried: list[tuple[Document, Job]] = []
        skipped_job_ids: list[str] = []
        candidate_jobs: list[Job] = []
        if job_ids:
            normalized_job_ids = list(
                dict.fromkeys(
                    str(job_id).strip() for job_id in job_ids if str(job_id).strip()
                )
            )[:limit]
            if not normalized_job_ids:
                return [], []
            candidate_jobs = list(
                self.db.scalars(
                    select(Job).where(Job.job_id.in_(normalized_job_ids))
                ).all()
            )
            existing_job_ids = {job.job_id for job in candidate_jobs}
            skipped_job_ids.extend(
                job_id for job_id in normalized_job_ids if job_id not in existing_job_ids
            )
        else:
            candidate_jobs = list(
                self.db.scalars(
                    select(Job)
                    .where(
                        Job.job_type == "document_ingestion",
                        Job.status == "failed",
                    )
                    .order_by(Job.updated_at.asc())
                    .limit(limit)
                ).all()
            )

        for job in candidate_jobs[:limit]:
            try:
                retried.append(self._retry_document_job(job))
            except DocumentIngestionStateError:
                skipped_job_ids.append(job.job_id)
        return retried, skipped_job_ids

    def _retry_document_job(self, job: Job) -> tuple[Document, Job]:
        if job.job_type != "document_ingestion":
            raise DocumentIngestionStateError("当前任务类型不支持重试。", status_code=409)
        if job.status != "failed":
            raise DocumentIngestionStateError("只有失败任务才允许重试。", status_code=409)

        document = self.db.scalar(
            select(Document).where(Document.document_id == job.target_id)
        )
        if document is None:
            raise DocumentIngestionStateError("任务对应的文档不存在。", status_code=404)

        file_path = Path(document.file_path)
        if not file_path.exists():
            raise DocumentIngestionStateError("源文件不存在，无法重试。", status_code=409)

        job.status = "pending"
        job.progress = 0.0
        job.error_message = ""
        document.status = "processing"
        self.db.add(job)
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        self.db.refresh(job)
        return document, job


class DocumentIngestionStateError(RuntimeError):
    def __init__(self, detail: str, *, status_code: int) -> None:
        super().__init__(detail)
        self.status_code = status_code


def process_document_ingestion_job(document_id: str, job_id: str) -> None:
    with SessionLocal() as db:
        document = db.scalar(
            select(Document).where(Document.document_id == document_id)
        )
        job = db.scalar(select(Job).where(Job.job_id == job_id))
        if not document or not job:
            return

        try:
            job.status = "running"
            job.progress = 10.0
            document.status = "processing"
            db.commit()

            extracted_text = _extract_text(Path(document.file_path))
            if not extracted_text.strip():
                raise ValueError("当前文件没有可解析的文本内容。")

            job.progress = 35.0
            db.commit()

            ingested_chunks = run_document_ingestion_pipeline(
                text=extracted_text,
                document_id=document.document_id,
                knowledge_base_id=document.knowledge_base_id,
                file_name=document.file_name,
                file_path=document.file_path,
                mime_type=document.mime_type,
            )
            if not ingested_chunks:
                raise ValueError("当前文件分块结果为空。")

            existing_chunk_ids = list(
                db.scalars(
                    select(DocumentChunk.chunk_id).where(
                        DocumentChunk.document_id == document.document_id
                    )
                ).all()
            )
            db.execute(
                delete(DocumentChunk).where(DocumentChunk.document_id == document.document_id)
            )
            chunk_rows = []
            vector_rows = []
            for chunk in ingested_chunks:
                chunk_id = str(uuid4())
                chunk_rows.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        document_id=document.document_id,
                        knowledge_base_id=document.knowledge_base_id,
                        chunk_index=chunk.chunk_index,
                        content=chunk.content,
                        char_count=chunk.char_count,
                    )
                )
                vector_row = {
                    "chunk_id": chunk_id,
                    "document_id": document.document_id,
                    "knowledge_base_id": document.knowledge_base_id,
                    "chunk_index": chunk.chunk_index,
                    "file_name": document.file_name,
                    "content": chunk.content,
                    "char_count": chunk.char_count,
                }
                vector_row.update(chunk.metadata)
                vector_rows.append(
                    vector_row
                )
            db.add_all(chunk_rows)
            job.progress = 60.0
            db.commit()

            qdrant_store = QdrantChunkStore()
            qdrant_store.delete_chunk_ids(existing_chunk_ids)
            qdrant_store.upsert_chunks(vector_rows)
            job.progress = 90.0
            db.commit()

            document.status = "ready"
            job.status = "completed"
            job.progress = 100.0
            job.error_message = ""
            db.commit()
        except Exception as exc:  # pragma: no cover
            document.status = "failed"
            job.status = "failed"
            job.error_message = str(exc)
            db.commit()


def _extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_path)
    if suffix == ".docx":
        return _extract_docx_text(file_path)
    if suffix == ".doc":
        return _extract_doc_text(file_path)

    raw = file_path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1", errors="ignore")


def validate_supported_document_filename(file_name: str) -> None:
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentIngestionStateError(
            f"暂不支持该文件格式，请上传：{SUPPORTED_DOCUMENT_FORMAT_LABEL}。",
            status_code=415,
        )


def _extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(text.strip() for text in page_texts if text.strip())


def _extract_docx_text(file_path: Path) -> str:
    document = DocxDocument(str(file_path))
    blocks: list[str] = []
    blocks.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append("\t".join(cells))
    return "\n\n".join(blocks)


def _extract_doc_text(file_path: Path) -> str:
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except FileNotFoundError as exc:
        raise ValueError("解析 .doc 文件需要安装 antiword。") from exc
    except subprocess.TimeoutExpired as exc:
        raise ValueError("解析 .doc 文件超时。") from exc
    except subprocess.CalledProcessError as exc:
        error_text = (exc.stderr or exc.stdout or "").strip()
        message = f"解析 .doc 文件失败：{error_text}" if error_text else "解析 .doc 文件失败。"
        raise ValueError(message) from exc
    return result.stdout
