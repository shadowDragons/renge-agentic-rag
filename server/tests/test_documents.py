from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.db.session import SessionLocal
from app.models import Document, Job
from app.services.document_ingestion import _extract_text


def test_upload_document_and_list_jobs(client: TestClient) -> None:
    upload_response = client.post(
        "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
        files={
            "file": (
                "测试制度.md",
                "# 测试制度\n\n这里是一段测试内容。".encode("utf-8"),
                "text/markdown",
            )
        },
    )
    assert upload_response.status_code == 201
    upload_data = upload_response.json()
    assert upload_data["document"]["knowledge_base_id"] == "kb-demo-001"
    assert upload_data["document"]["file_name"] == "测试制度.md"
    assert upload_data["job"]["job_type"] == "document_ingestion"

    document_response = client.get("/api/v1/knowledge-bases/kb-demo-001/documents")
    assert document_response.status_code == 200
    documents = document_response.json()
    assert any(item["file_name"] == "测试制度.md" for item in documents)

    jobs_response = client.get("/api/v1/jobs", params={"job_type": "document_ingestion"})
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    job_item = next(item for item in jobs if item["job_id"] == upload_data["job"]["job_id"])
    assert job_item["target_type"] == "document"
    assert job_item["target_name"] == "测试制度.md"
    assert job_item["knowledge_base_id"] == "kb-demo-001"
    assert job_item["knowledge_base_name"] == "默认知识库"
    assert job_item["target_status"] in {"processing", "ready", "failed"}
    assert job_item["sla"]["policy_key"] == "document_ingestion"
    assert job_item["sla"]["status"] in {"normal", "completed", "failed"}


def test_upload_accepts_legacy_doc_file(client: TestClient) -> None:
    upload_response = client.post(
        "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
        files={
            "file": (
                "旧版制度.doc",
                b"legacy-doc-content",
                "application/msword",
            )
        },
    )
    assert upload_response.status_code == 201
    assert upload_response.json()["document"]["file_name"] == "旧版制度.doc"


def test_extract_docx_text_includes_paragraphs_and_tables(tmp_path: Path) -> None:
    docx_path = tmp_path / "制度.docx"
    docx = DocxDocument()
    docx.add_paragraph("第一章 总则")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    docx.save(docx_path)

    text = _extract_text(docx_path)

    assert "第一章 总则" in text
    assert "字段\t说明" in text


def test_delete_document_removes_chunks_and_jobs(client: TestClient) -> None:
    upload_response = client.post(
        "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
        files={
            "file": (
                "待删除文档.md",
                "# 待删除文档\n\n这里是一段会被删除的内容。".encode("utf-8"),
                "text/markdown",
            )
        },
    )
    assert upload_response.status_code == 201
    upload_data = upload_response.json()
    document_id = upload_data["document"]["document_id"]
    job_id = upload_data["job"]["job_id"]

    delete_response = client.delete(
        f"/api/v1/knowledge-bases/kb-demo-001/documents/{document_id}"
    )
    assert delete_response.status_code == 200
    delete_data = delete_response.json()
    assert delete_data["document_id"] == document_id
    assert delete_data["deleted_job_count"] >= 1

    document_response = client.get("/api/v1/knowledge-bases/kb-demo-001/documents")
    assert document_response.status_code == 200
    documents = document_response.json()
    assert all(item["document_id"] != document_id for item in documents)

    jobs_response = client.get("/api/v1/jobs", params={"job_type": "document_ingestion"})
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    assert all(item["job_id"] != job_id for item in jobs)


def test_retry_failed_document_job(client: TestClient) -> None:
    upload_response = client.post(
        "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
        files={
            "file": (
                "失败文档.md",
                "   \n\t".encode("utf-8"),
                "text/markdown",
            )
        },
    )
    assert upload_response.status_code == 201
    upload_data = upload_response.json()
    document_id = upload_data["document"]["document_id"]
    job_id = upload_data["job"]["job_id"]

    failed_job_response = client.get(f"/api/v1/jobs/{job_id}")
    assert failed_job_response.status_code == 200
    failed_job = failed_job_response.json()
    assert failed_job["status"] == "failed"
    assert failed_job["retryable"] is True

    with SessionLocal() as db:
        document = db.get(Document, document_id)
        assert document is not None
        with open(document.file_path, "w", encoding="utf-8") as handle:
            handle.write("# 修复后文档\n\n现在有可解析的文本内容。")

    retry_response = client.post(f"/api/v1/jobs/{job_id}/retry")
    assert retry_response.status_code == 200
    retry_data = retry_response.json()
    assert retry_data["job_id"] == job_id
    assert retry_data["retryable"] is False

    retried_job_response = client.get(f"/api/v1/jobs/{job_id}")
    assert retried_job_response.status_code == 200
    retried_job = retried_job_response.json()
    assert retried_job["status"] == "completed"
    assert retried_job["target_status"] == "ready"
    assert retried_job["retryable"] is False
    assert retried_job["sla"]["status"] == "completed"


def test_list_jobs_supports_sla_status_filter(client: TestClient) -> None:
    upload_response = client.post(
        "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
        files={
            "file": (
                "超时文档.md",
                "# 超时文档\n\n这里是一段测试内容。".encode("utf-8"),
                "text/markdown",
            )
        },
    )
    assert upload_response.status_code == 201
    job_id = upload_response.json()["job"]["job_id"]

    with SessionLocal() as db:
        job = db.get(Job, job_id)
        assert job is not None
        stale_time = datetime.now(timezone.utc) - timedelta(minutes=10)
        job.status = "running"
        job.progress = 45
        job.created_at = stale_time
        job.updated_at = stale_time
        db.add(job)
        db.commit()

    jobs_response = client.get(
        "/api/v1/jobs",
        params={"job_type": "document_ingestion", "sla_status": "breached"},
    )
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    job_item = next(item for item in jobs if item["job_id"] == job_id)
    assert job_item["sla"]["status"] == "breached"
    assert job_item["sla"]["breach_seconds"] >= 1


def test_retry_failed_document_jobs_in_batch(client: TestClient) -> None:
    failed_job_ids: list[str] = []
    document_ids: list[str] = []

    for index in range(2):
        upload_response = client.post(
            "/api/v1/knowledge-bases/kb-demo-001/documents/upload",
            files={
                "file": (
                    f"失败文档-{index}.md",
                    "   \n\t".encode("utf-8"),
                    "text/markdown",
                )
            },
        )
        assert upload_response.status_code == 201
        payload = upload_response.json()
        failed_job_ids.append(payload["job"]["job_id"])
        document_ids.append(payload["document"]["document_id"])

    with SessionLocal() as db:
        for document_id in document_ids:
            document = db.get(Document, document_id)
            assert document is not None
            with open(document.file_path, "w", encoding="utf-8") as handle:
                handle.write(f"# 修复后文档 {document_id}\n\n现在有可解析的文本内容。")

    retry_response = client.post(
        "/api/v1/jobs/retry-batch",
        json={"job_ids": failed_job_ids, "limit": 2},
    )
    assert retry_response.status_code == 200
    retry_data = retry_response.json()
    assert retry_data["requested_count"] == 2
    assert retry_data["retried_count"] == 2
    assert retry_data["skipped_count"] == 0

    jobs_response = client.get("/api/v1/jobs", params={"job_type": "document_ingestion"})
    assert jobs_response.status_code == 200
    jobs = jobs_response.json()
    retried_jobs = [item for item in jobs if item["job_id"] in failed_job_ids]
    assert len(retried_jobs) == 2
    assert all(item["status"] == "completed" for item in retried_jobs)
    assert all(item["target_status"] == "ready" for item in retried_jobs)
